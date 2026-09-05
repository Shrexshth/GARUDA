import os
import uuid
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

def generate_nfa_docx(payload_data: dict) -> str:
    """
    Takes a payload dictionary, injects it into a docx template, and saves to the outputs directory.
    Uses docxtpl to render data into MRPL format templates.
    """
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        return "Error: docxtpl is not installed. Please install it."

    # Fetch configuration from env
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    templates_dir_env = os.getenv("TEMPLATES_DIR")
    outputs_dir_env = os.getenv("OUTPUTS_DIR")

    template_dir = templates_dir_env if templates_dir_env else os.path.join(base_dir, "templates")
    output_dir = outputs_dir_env if outputs_dir_env else os.path.join(base_dir, "data", "outputs")
    
    template_path = os.path.join(template_dir, "nfa_template.docx")
    
    # Safely create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate professional output filename
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    output_filename = f"NFA_Draft_{timestamp}.docx"
    output_path = os.path.join(output_dir, output_filename)
    
    try:
        # Check if template exists, if not, we create a blank template dynamically (for robustness)
        if not os.path.exists(template_path):
            os.makedirs(os.path.dirname(template_path), exist_ok=True)
            try:
                from docx import Document
            except ImportError:
                return "Error: python-docx is not installed."
                
            doc = Document()
            doc.add_heading('Note for Approval', 0)
            doc.add_paragraph('{{ summary }}')
            doc.add_paragraph('Status: {{ status }}')
            doc.save(template_path)
            
        doc = DocxTemplate(template_path)
        doc.render(payload_data)
        doc.save(output_path)
        
        return output_path
    except Exception as e:
        return f"Error generating document: {str(e)}"
